# -*- coding: utf-8 -*-
"""Tests for DOCX extractor — table inclusion fix."""
import sys
import os
import tempfile
from pathlib import Path

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

import docx as _docx
import pytest

from uploaded_doc.extractor import extract_docx


def _make_docx(tmp_dir: Path) -> Path:
    """Create a DOCX with a paragraph, a 3-row×2-col table, and a closing paragraph."""
    doc = _docx.Document()
    doc.add_paragraph("Uvodni paragraf ugovora.")

    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Stavka"
    table.cell(0, 1).text = "Vrednost"
    table.cell(1, 0).text = "Zarada"
    table.cell(1, 1).text = "80000 RSD"
    table.cell(2, 0).text = "Radno vreme"
    table.cell(2, 1).text = "40 sati nedeljno"

    doc.add_paragraph("Završni paragraf.")

    path = tmp_dir / "test_ugovor.docx"
    doc.save(str(path))
    return path


def test_extract_docx_includes_table_cells(tmp_path):
    """All table cell values must appear in extracted text."""
    path = _make_docx(tmp_path)
    text, is_scanned, _ = extract_docx(path)

    assert "80000 RSD" in text, f"Table cell '80000 RSD' missing. Got: {text[:400]}"
    assert "Radno vreme" in text, f"Table cell 'Radno vreme' missing."
    assert "40 sati nedeljno" in text, f"Table cell '40 sati nedeljno' missing."
    assert "Stavka" in text
    assert "Vrednost" in text
    assert is_scanned is False


def test_extract_docx_includes_paragraphs(tmp_path):
    """Paragraph text must still be present alongside table content."""
    path = _make_docx(tmp_path)
    text, _, __ = extract_docx(path)

    assert "Uvodni paragraf ugovora." in text
    assert "Završni paragraf." in text


def test_extract_docx_preserves_order(tmp_path):
    """Paragraph must appear before table rows in the output."""
    path = _make_docx(tmp_path)
    text, _, __ = extract_docx(path)

    para_pos = text.find("Uvodni paragraf")
    table_pos = text.find("80000 RSD")
    closing_pos = text.find("Završni paragraf")

    assert para_pos < table_pos, "Opening paragraph must come before table data"
    assert table_pos < closing_pos, "Table data must come before closing paragraph"


def test_extract_docx_no_table_no_regression(tmp_path):
    """Documents without tables must still work correctly."""
    doc = _docx.Document()
    doc.add_paragraph("Samo tekst bez tabela.")
    doc.add_paragraph("Drugi paragraf.")
    path = tmp_path / "no_table.docx"
    doc.save(str(path))

    text, is_scanned, _ = extract_docx(path)
    assert "Samo tekst bez tabela." in text
    assert "Drugi paragraf." in text
    assert is_scanned is False
