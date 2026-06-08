# -*- coding: utf-8 -*-
"""
F6 — DOCX export: konvertuje tekst analize / nacrta / strategije u formatiran .docx fajl.
"""
from __future__ import annotations

import re
from datetime import datetime
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

_BLUE   = RGBColor(0x1A, 0x5F, 0x7A)
_DARK   = RGBColor(0x0D, 0x1B, 0x2A)
_GOLD   = RGBColor(0xC9, 0xA8, 0x4C)
_GREY   = RGBColor(0x88, 0x88, 0x88)

_SEKCIJA_RE = re.compile(
    r"^(\d+\.\s+[A-ZŠĐŽČĆА-Я\s]{4,}|[A-ZŠĐŽČĆ\s]{4,}:)$"
)
_KLJUC_KW = ("PREPORUKA", "ZAKLJUČAK", "UKUPNA OCENA", "PRELIMINARNI STAV", "PREPORUČENA AKCIJA")


def tekst_u_docx(naslov: str, tekst: str, tip: str = "analiza") -> bytes:
    """Konvertuje tekst u formatiran .docx. Vraća bytes za download."""
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # Header — Vindex AI + datum
    hdr = doc.add_heading("", level=0)
    r = hdr.add_run("VINDEX AI")
    r.font.size = Pt(11)
    r.font.color.rgb = _BLUE
    r.font.bold = True
    hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    dr = hdr.add_run(f"\n{datetime.now().strftime('%d.%m.%Y')}")
    dr.font.size = Pt(9)
    dr.font.color.rgb = _GREY

    # Naslov dokumenta
    t = doc.add_heading(naslov[:200], level=1)
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in t.runs:
        run.font.color.rgb = _DARK

    doc.add_paragraph("")

    # Parsiraj i formatiraj liniju po liniju
    for linija in tekst.split("\n"):
        s = linija.strip()
        if not s:
            doc.add_paragraph("")
            continue

        if _SEKCIJA_RE.match(s):
            p = doc.add_heading(s, level=2)
            for run in p.runs:
                run.font.color.rgb = _BLUE

        elif s.startswith(("-", "•", "–", "*")):
            p = doc.add_paragraph(s.lstrip("-•–* "), style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.3)

        elif any(kw in s.upper() for kw in _KLJUC_KW):
            p = doc.add_paragraph()
            run = p.add_run(s)
            run.font.bold = True
            run.font.color.rgb = _GOLD

        else:
            doc.add_paragraph(s)

    # Footer
    doc.add_paragraph("")
    fp = doc.add_paragraph()
    fr = fp.add_run(
        "Generisano uz pomoć Vindex AI (vindex.rs) — za informativne svrhe. "
        "Ne predstavlja pravni savet."
    )
    fr.font.size = Pt(8)
    fr.font.color.rgb = _GREY
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
