# -*- coding: utf-8 -*-
"""
Vindex AI — routers/drafting.py

Nacrti, analiza, sazmi, feedback, podnesci, AI paralegal pipeline

Core Consolidation Sec 1.4 (2026-07-22) — interim ownership (NOT merged,
pilot-gated by explicit founder decision): this file hosts TWO of the
three drafting surfaces. POST /api/nacrt (line ~236) wraps
drafting.router.generate_draft (drafting/router.py) unchanged — the quick
single-shot path. POST /api/podnesak (line ~372) is this file's OWN
generation logic against templates/podnesci.py's TIPOVI/SABLONI, RAG-
augmented (sudska praksa retrieval, see rag_upit below) and tied to an
open predmet — richer context than /api/nacrt. 6 of podnesci.py's 12
types duplicate drafting/templates.py's TEMPLATES under the same type
keys with different template text (tuzba_naknada_stete, tuzba_radni_spor,
tuzba_razvod, prigovor_platni_nalog, krivicna_prijava,
predlog_privremena_mera) — known, tracked duplication, not an oversight.
Do not resolve by intuition; see VINDEX_CORE_CONSOLIDATION.md Sec 1.4.
"""
import asyncio
import json
import logging
import os
import time as _time
from typing import Optional, List

from fastapi import APIRouter, Depends, HTTPException, Request, UploadFile, File
from fastapi.responses import JSONResponse, Response
from pydantic import BaseModel, Field, field_validator

from security.html_sanitize import sanitize_user_input
from shared.deps import _audit, _get_supa, _q_hash, get_current_user
from shared.rate import limiter
from shared.permissions import PermissionService
from shared.usage import UsageService
from shared.sentry import capture_exception as _sentry_capture
from shared.llm_retry import llm_retry

from main import ask_analiza, _skini_pii
from drafting.router import generate_draft as _drafting_generate
from drafting.templates import get_types_list as _drafting_get_types
from app.services import audit_log as _al
from templates.podnesci import (
    TIPOVI as PODNESAK_TIPOVI,
    EKSTRAKCIONI_PROMPTOVI,
    OBOGACIVANJE_PROMPTOVI,
    popuni_sablon,
)
from knowledge.vks_standards import preporuci_iznose as vks_preporuci

logger = logging.getLogger("vindex.api")
router = APIRouter()

_MAX_PLAYBOOK_BYTES = 2 * 1024 * 1024  # 2 MB


@llm_retry
def _pozovi_drafting_api(client, **kwargs):
    """Zajednički retry-zaštićen poziv za sync OpenAI pozive u ovom fajlu
    (sazmi, podnesak ekstrakcija/obogaćivanje).

    CELINA 4 (2026-07-24): @llm_retry -- max 3 pokušaja sa exponential
    backoff-om za rate-limit/5xx/timeout/connection greške.
    """
    return client.chat.completions.create(**kwargs)


# ── Models ────────────────────────────────────────────────────────────────────

class NacrtReq(BaseModel):
    vrsta: str = Field(..., min_length=2, max_length=200)
    opis:  str = Field(..., min_length=10, max_length=5000)
    # Institutional Memory V2 (2026-07-26): opciono -- ako je prosleđen,
    # generisani nacrt ide u staging_memory na advokatsku overu (v.
    # _stage_draft_for_review niže) umesto da se odmah zaboravi. Bez ovoga
    # ponašanje je nepromenjeno (samo generiši i vrati tekst).
    predmet_id: Optional[str] = Field(None, max_length=100)

    @field_validator("vrsta", "opis")
    @classmethod
    def ocisti(cls, v: str) -> str:
        return sanitize_user_input(v.strip()) or ""


class AnalizaReq(BaseModel):
    tekst:   str = Field(..., min_length=10, max_length=50000)
    pitanje: str = Field("", max_length=1000)

    @field_validator("tekst", "pitanje")
    @classmethod
    def ocisti(cls, v: str) -> str:
        return sanitize_user_input(v.strip()) or ""


class FeedbackReq(BaseModel):
    pitanje: str = Field("", max_length=2000)
    odgovor: str = Field("", max_length=5000)
    tip:     str = Field("greska", max_length=50)

    @field_validator("pitanje", "odgovor")
    @classmethod
    def _sanitize(cls, v: str) -> str:
        return sanitize_user_input(v) or ""


class PodnesakReq(BaseModel):
    tip:        str           = Field(..., max_length=50)
    opis:       str           = Field(..., min_length=20, max_length=5000)
    sud_naziv:  Optional[str] = Field(None, max_length=200)
    sud_adresa: Optional[str] = Field(None, max_length=300)
    # Institutional Learning & RAG Audit (2026-07-26) #3 — v. NacrtReq.predmet_id.
    predmet_id: Optional[str] = Field(None, max_length=100)

    @field_validator("tip")
    @classmethod
    def validiraj_tip(cls, v: str) -> str:
        dozvoljeni = {
            "tuzba_naknada_stete", "zalba_parnicna", "predlog_izvrsenje",
            "tuzba_radni_spor", "tuzba_razvod", "prigovor_platni_nalog",
            "krivicna_prijava", "predlog_privremena_mera",
            "odgovor_na_tuzbu", "zalba_krivicna", "urgencija_sudu", "prigovor_izvrsenje",
        }
        if v not in dozvoljeni:
            raise ValueError(f"Tip podneska mora biti jedan od: {dozvoljeni}")
        return v

    @field_validator("opis")
    @classmethod
    def ocisti_opis(cls, v: str) -> str:
        return sanitize_user_input(v.strip()) or ""

    @field_validator("sud_naziv", "sud_adresa")
    @classmethod
    def ocisti_sud(cls, v: Optional[str]) -> Optional[str]:
        return sanitize_user_input(v.strip()) if v else None


class NacrtChecklistReq(BaseModel):
    tip:       str = Field(..., max_length=60)
    cinjenice: str = Field(..., min_length=10, max_length=8000)

    @field_validator("tip")
    @classmethod
    def validiraj_tip(cls, v: str) -> str:
        from nacrti.checklist_config import SVI_TIPOVI
        if v not in SVI_TIPOVI:
            raise ValueError(f"Nepoznat tip podneska: {v!r}. Dozvoljeni: {SVI_TIPOVI}")
        return v

    @field_validator("cinjenice")
    @classmethod
    def ocisti_cinjenice(cls, v: str) -> str:
        return sanitize_user_input(v.strip()) or ""


class SazmiReq(BaseModel):
    odgovor: str = Field(..., max_length=6000)
    format:  str = Field(default="email", pattern="^(email|viber|pisano)$")


# ── Helpers ───────────────────────────────────────────────────────────────────

async def _pokreni(fn, *args):
    return await asyncio.to_thread(fn, *args)


def _normalizuj_rezultat(rezultat: dict, credits_remaining: Optional[int] = None) -> dict:
    resp: dict = {}
    if not isinstance(rezultat, dict):
        resp["odgovor"] = str(rezultat)
    elif rezultat.get("status") == "success":
        resp["odgovor"] = rezultat.get("data", "")
    else:
        resp["odgovor"] = rezultat.get(
            "message",
            "Došlo je do greške prilikom obrade zahteva. Pokušajte ponovo.",
        )
    if credits_remaining is not None:
        resp["credits_remaining"] = credits_remaining
    return resp


def _greska_odgovor(status_code: int, poruka: str) -> JSONResponse:
    logger.warning("API greška %d: %s", status_code, poruka)
    return JSONResponse(status_code=status_code, content={"greska": poruka})


# Institutional Memory Architecture V2 (2026-07-26), STUB 1 — Quality Gate &
# Staging Memory. ZAMENJUJE raniju _index_finalized_draft (Institutional
# Learning & RAG Audit, 2026-07-25), koja je AI-generisan nacrt indeksirala
# DIREKTNO u kancelarija_{id}/user_{id} Pinecone namespace čim je predmet_id
# prosleđen -- to je tačno "toksično učenje" rizik koji V2 sprečava: nikad
# više se sirov, neproveren AI tekst ne upisuje u glavnu bazu znanja.
# Sada: finalizovan nacrt ide u staging_memory (migracija 088) sa
# automatskim Quality Gate skorom (services/quality_gate.py) i čeka
# eksplicitnu advokatsku potvrdu (v. POST /api/staging/{id}/approve) PRE
# nego što ijedan vektor uđe u Pinecone.
async def _stage_draft_for_review(user: dict, predmet_id: str, tip: str, naziv: str, tekst: str) -> None:
    """Fire-and-forget (asyncio.create_task) -- ne sme da uspori odgovor
    korisniku niti da greška ovde ikad utiče na već generisan i vraćen nacrt."""
    try:
        supa = _get_supa()
        pred_row = await asyncio.to_thread(
            lambda: supa.table("predmeti").select("id").eq("id", predmet_id).eq("user_id", user["user_id"]).maybe_single().execute()
        )
        if not pred_row.data:
            logger.warning("[STAGING] predmet_id=%s ne pripada korisniku — preskočeno", predmet_id)
            return

        from services.quality_gate import evaluate_draft_quality
        from shared.kancelarija_utils import get_kancelarija_id

        quality = await evaluate_draft_quality(tekst, tip)
        kancelarija_id = await get_kancelarija_id(supa, user["user_id"])

        await asyncio.to_thread(
            lambda: supa.table("staging_memory").insert({
                "user_id": user["user_id"],
                "kancelarija_id": kancelarija_id,
                "predmet_id": predmet_id,
                "tip": tip,
                "naziv": naziv,
                "tekst": tekst,
                "confidence_score": quality["confidence_score"],
                "quality_detail": quality["detail"],
            }).execute()
        )
        logger.info(
            "[STAGING] predmet=%s tip=%s confidence=%.2f -- čeka advokatsku potvrdu",
            predmet_id, tip, quality["confidence_score"],
        )
    except Exception as exc:
        _sentry_capture(exc)
        logger.warning("[STAGING] neuspešno (non-fatal) predmet=%s: %s", predmet_id, exc)


_APPROVAL_CONFIDENCE_THRESHOLD = 0.85


async def _promote_staged_draft_to_pinecone(supa, staging_row: dict) -> bool:
    """Poziva se SAMO iz POST /api/staging/{id}/approve, i SAMO ako
    is_lawyer_approved=True I confidence_score >= 0.85 (oba uslova, v.
    migracija 088). Indeksira pod origin=LAWYER_VERIFIED (ne AI_GENERATED --
    upravo je prošao ljudsku overu), sa parent_id/origin_chain koji čuvaju
    da je tekst NASTAO kao AI_GENERATED pre nego što je postao
    LAWYER_VERIFIED (STUB 2 -- lineage protiv AI-to-AI degeneracije: budući
    retrieval uvek zna da je ovo poreklo bilo AI, ne izvorni advokatski rad
    od nule). Vraća True ako je Pinecone upis uspeo."""
    from uploaded_doc.chunker import chunk_document
    from uploaded_doc.ingest import ingest_session
    from uploaded_doc.session import generate_session_id
    from shared.kancelarija_utils import rag_owner_namespace
    from shared.vector_origin import ORIGIN_LAWYER_VERIFIED, ORIGIN_AI_GENERATED, now_iso

    predmet_id = staging_row["predmet_id"]
    user_id = staging_row["user_id"]
    kancelarija_id = staging_row.get("kancelarija_id")
    tekst = staging_row["tekst"]
    naziv = staging_row.get("naziv") or staging_row.get("tip") or "Nacrt"

    source_meta = {
        "source_filename": f"Nacrt — {naziv}", "source_format": "txt",
        "source_sha256": "", "is_scanned": False, "session_id": "__local__",
    }
    manifest = await asyncio.to_thread(chunk_document, tekst, source_meta)
    if manifest.total_chunks == 0:
        return False

    owner_ns = rag_owner_namespace(user_id, kancelarija_id)
    session_id = generate_session_id()

    try:
        await asyncio.to_thread(
            ingest_session, manifest, session_id,
            namespace_override=owner_ns,
            extra_metadata={
                "predmet_id": predmet_id,
                "kancelarija_id": kancelarija_id or "",
                "type": "draft_final",
                "origin": ORIGIN_LAWYER_VERIFIED,
                "parent_id": staging_row["id"],
                "origin_chain": [ORIGIN_AI_GENERATED, ORIGIN_LAWYER_VERIFIED],
                "created_at": now_iso(),
                "golden_template": False,
            },
        )
    except Exception as pe:
        logger.warning("[STAGING_PROMOTE] Pinecone ingest neuspešan predmet=%s: %s", predmet_id, str(pe)[:150])
        return False

    try:
        _rn_res = await asyncio.to_thread(
            lambda: supa.table("predmet_dokumenti").select("redni_broj").eq("predmet_id", predmet_id).order("redni_broj", desc=True).limit(1).execute()
        )
        next_rn = int((_rn_res.data or [{}])[0].get("redni_broj") or 0) + 1
    except Exception:
        next_rn = 1

    await asyncio.to_thread(
        lambda: supa.table("predmet_dokumenti").insert({
            "predmet_id": predmet_id, "user_id": user_id,
            "naziv_fajla": f"AI Nacrt (odobren) — {naziv}",
            "storage_path": f"draft/{session_id}",
            "pinecone_namespace": owner_ns, "status": "indeksirano",
            "velicina_kb": max(1, len(tekst.encode("utf-8")) // 1024),
            "redni_broj": next_rn, "tekst_sadrzaj": tekst[:100_000],
        }).execute()
    )
    return True


def _izvori_kontekst(docs: list[str], limit: int = 4) -> str:
    """FAZA 3 (2026-07-24): označava svaki RAG dokument kao [IZVOR-n] -- isti
    identity-based mapping obrazac kao services/legal_reasoning_engine.py
    (SOURCE-n). Omogućava sistemskom promptu da traži eksplicitno pozivanje
    na [IZVOR-n] pri navođenju člana/prakse, i critique pass-u da proveri
    da li je svaki navod stvarno potkrepljen tim izvorom."""
    if not docs:
        return ""
    return "\n\n".join(f"[IZVOR-{i + 1}]\n{d}" for i, d in enumerate(docs[:limit]))


# ── FAZA 3 (2026-07-24): Critique & Self-Correction pass za podneske ──────────

_CRITIQUE_SYSTEM = """\
Ti si strogi glavni advokat (senior partner) koji vrši finalnu proveru nacrta \
pravnog podneska pre nego što ide advokatu na potpis. Ne pišeš nacrt iznova — \
proveravaš ga i ispravljaš SAMO ako postoji stvaran problem.

TVOJ ZADATAK IMA DVA DELA:

1. PROVERA HALUCINACIJA:
Zakonski kontekst dostavljen uz nacrt je označen kao [IZVOR-1], [IZVOR-2], itd.
Proveri da li nacrt navodi članove zakona, brojeve presuda ili druge pravne
reference koje NISU potkrepljene tim izvorima. Osnovne, opštepoznate odredbe
(npr. ZOO čl. 154, ZPP čl. 194) NISU halucinacija ni bez izvora — problem je
SAMO kad je naveden konkretan, neuobičajen ili sporan broj člana/presude koji
se ne može potvrditi ni iz izvora ni iz opšteg pravnog znanja.

2. PROVERA OBAVEZNIH FORMALNIH ELEMENATA:
Proveri da nacrt sadrži: (a) zaglavlje sa nazivom/adresom suda, (b) jasnu
identifikaciju stranaka, (c) činjenično stanje, (d) pravni osnov sa referencama
na zakon, (e) dokazni predlog, (f) tužbeni zahtev/petitum (ako je tip podneska
takav da petitum zahteva) koji je izvršan i nedvosmislen, (g) mesto/datum i
potpisni blok.

Vrati ISKLJUČIVO validan JSON, bez markdown blokova, bez teksta van JSON-a:
{
  "ima_izmisljenih_navoda": true/false,
  "izmisljeni_navodi": ["kratak opis svakog spornog navoda"],
  "nedostaju_elementi": ["naziv elementa koji nedostaje ili je nepotpun"],
  "ispravljen_tekst": "CEO ispravljen tekst nacrta. Ako NEMA problema (oba polja iznad su prazna/false), vrati IDENTIČAN originalni tekst bez ijedne izmene."
}

STROGA PRAVILA:
- Ne diraj stilske izbore koji nisu greška — ispravljaj SAMO stvarne probleme
  (halucinacije, nedostajući obavezni elementi, očigledne pravne greške).
- Ako ukloniš izmišljen član zakona, zameni ga generičkom formulacijom bez
  izmišljenog broja ("u skladu sa važećim propisima") ili placeholder-om
  "[proveriti relevantan član]" — NIKAD ne izmišljaj zamenski broj.
- "ispravljen_tekst" mora biti KOMPLETAN nacrt od početka do kraja, ne samo
  izmenjeni delovi ili rezime.
"""


@llm_retry
async def _pozovi_kriticara(oai_client, nacrt: str, kontekst: str, tip_naziv: str) -> dict:
    """Retry-ovani deo critique pass-a -- izdvojen kao poseban helper (isti
    razlog kao services/legal_reasoning_engine.py::_pozovi_reasoning_api):
    _critique_and_refine_draft ima sopstveni try/except sa fallback-om na
    originalni nacrt, koji bi inače sakrio grešku od tenacity-ja."""
    resp = await asyncio.wait_for(
        asyncio.to_thread(
            lambda: oai_client.chat.completions.create(
                model="gpt-4o-mini",
                temperature=0,
                max_tokens=4000,
                response_format={"type": "json_object"},
                messages=[
                    {"role": "system", "content": _CRITIQUE_SYSTEM},
                    {"role": "user", "content": (
                        f"VRSTA PODNESKA: {tip_naziv}\n\n"
                        f"IZVORI (RAG kontekst sa oznakama):\n"
                        f"{kontekst or '(nema dostavljenih izvora -- oslanjaj se samo na opšte pravno znanje)'}\n\n"
                        f"NACRT ZA PROVERU:\n{nacrt}"
                    )},
                ],
            )
        ),
        timeout=30.0,
    )
    raw = (resp.choices[0].message.content or "{}").strip()
    return json.loads(raw)


async def _critique_and_refine_draft(nacrt: str, kontekst: str, tip: str, log_id: str) -> str:
    """Drugi, brz LLM prolaz nad već generisanim nacrtom: proverava izmišljene
    članove zakona van [IZVOR-n] konteksta i obavezne formalne elemente
    podneska, i ispravlja ih ako postoje. Ako je nacrt besprekoran, vraća ga
    NEIZMENJEN (ignoriše model's "ispravljen_tekst" kad model sam prijavi da
    nema problema -- izbegava suvišno prepisivanje/drift). Nikad ne baca:
    svaka greška pada nazad na originalni nacrt umesto da blokira odgovor."""
    try:
        from openai import OpenAI as _OAI
        oai_client = _OAI(api_key=os.getenv("OPENAI_API_KEY"))
        tip_naziv = PODNESAK_TIPOVI.get(tip, tip)
        kritika = await _pozovi_kriticara(oai_client, nacrt, kontekst, tip_naziv)

        izmisljeni = kritika.get("izmisljeni_navodi") or []
        nedostaje = kritika.get("nedostaju_elementi") or []
        ima_problema = bool(kritika.get("ima_izmisljenih_navoda")) or bool(izmisljeni) or bool(nedostaje)

        if not ima_problema:
            return nacrt

        ispravljen = (kritika.get("ispravljen_tekst") or "").strip()
        if not ispravljen:
            logger.warning(
                "Critique pass prijavio probleme ali nije vratio ispravljen tekst [q=%s]", log_id
            )
            return nacrt

        logger.info(
            "Critique pass ispravio nacrt [q=%s]: izmisljeni_navodi=%s nedostaju_elementi=%s",
            log_id, izmisljeni, nedostaje,
        )
        return ispravljen
    except Exception as exc:
        _sentry_capture(exc)
        logger.warning("Critique pass neuspešan [q=%s]: %s — vraćam originalni nacrt", log_id, exc)
        return nacrt


# ── Routes ────────────────────────────────────────────────────────────────────

@router.get("/api/nacrt/types")
async def nacrt_types():
    """Vraća listu dostupnih tipova nacrta (bez autentifikacije)."""
    return {"tipovi": _drafting_get_types()}


@router.get("/api/courts")
async def get_courts():
    """Katalog srpskih sudova sa adresama — za popunjavanje zaglavlja podnesaka."""
    from knowledge.sudovi import SUDOVI
    return {"sudovi": SUDOVI}


@router.post("/api/playbook/upload")
@limiter.limit("10/minute")
async def playbook_upload(
    request: Request,
    file: UploadFile = File(...),
    user: dict = Depends(PermissionService.require("drafting")),
):
    """P4.4 — Upload firm playbook (TXT or DOCX). Ne troši kredit."""
    from pathlib import Path as _Path
    import tempfile
    from uploaded_doc.extractor import DocumentSafetyLimitExceeded, extract_docx, extract_txt

    suffix = _Path(file.filename or "").suffix.lower()
    if suffix not in {".txt", ".docx"}:
        raise HTTPException(status_code=415, detail="Podržani formati: TXT, DOCX")

    raw = await file.read()
    if len(raw) > _MAX_PLAYBOOK_BYTES:
        raise HTTPException(status_code=413, detail="Fajl je preko 2MB")

    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmp.write(raw)
            tmp_path = _Path(tmp.name)

        try:
            if suffix == ".docx":
                tekst, _ = await asyncio.to_thread(extract_docx, tmp_path)
            else:
                tekst, _ = await asyncio.to_thread(extract_txt, tmp_path)
        except DocumentSafetyLimitExceeded:
            raise HTTPException(
                status_code=413,
                detail="Fajl je odbijen — sadržaj posle raspakivanja prelazi bezbednosni limit.",
            )

        if not tekst or not tekst.strip():
            raise HTTPException(status_code=422, detail="Fajl je prazan ili nečitljiv")

        from drafting.playbook import ingest_playbook
        count = await asyncio.to_thread(ingest_playbook, user["user_id"], file.filename or "", tekst)
        return {"filename": file.filename, "chunks_ingested": count}
    finally:
        if tmp_path and tmp_path.exists():
            try:
                tmp_path.unlink()
            except Exception as _exc:
                _sentry_capture(_exc)
                pass


@router.delete("/api/playbook")
@limiter.limit("10/minute")
async def playbook_delete(request: Request, user: dict = Depends(PermissionService.require("drafting"))):
    """P4.4 — Briše ceo playbook korisnika iz Pinecone."""
    from drafting.playbook import delete_playbook
    deleted = await asyncio.to_thread(delete_playbook, user["user_id"])
    return {"deleted_chunks": deleted}


@router.get("/api/playbook/status")
@limiter.limit("30/minute")
async def playbook_status(request: Request, user: dict = Depends(PermissionService.require("drafting"))):
    """P4.4 — Vraća status playbook-a: da li postoji i koliko chunks ima."""
    def _check():
        try:
            from uploaded_doc.ingest import _get_pinecone_index
            index = _get_pinecone_index()
            ns = f"playbook_{user['user_id']}"
            stats = index.describe_index_stats()
            ns_data = stats.namespaces.get(ns) if hasattr(stats, "namespaces") else None
            count = (ns_data.vector_count if hasattr(ns_data, "vector_count") else 0) if ns_data else 0
            return {"has_playbook": count > 0, "chunk_count": count}
        except Exception as _exc:
            _sentry_capture(_exc)
            return {"has_playbook": False, "chunk_count": 0}
    return await asyncio.to_thread(_check)


@router.post("/api/nacrt")
@limiter.limit("10/minute")
async def nacrt(req: NacrtReq, request: Request, user: dict = Depends(PermissionService.require("drafting"))):
    """Generisanje nacrta pravnog dokumenta (strukturirani šablon)."""
    logger.info("Nacrt [uid=%.8s] vrsta=%s", user["user_id"], req.vrsta)
    asyncio.create_task(_audit(user["user_id"], f"nacrt:{req.vrsta}", ""))
    try:
        qh_nacrt = _q_hash(_skini_pii(req.opis))
        t0 = _time.monotonic()
        rezultat = await _pokreni(_drafting_generate, req.vrsta, _skini_pii(req.opis), user["user_id"])
        latency_ms = int((_time.monotonic() - t0) * 1000)
        _al.log_response(
            endpoint="/api/nacrt",
            query_hash=qh_nacrt,
            tip=req.vrsta[:20],
            response_text=rezultat.get("data", ""),
            latency_ms=latency_ms,
        )
        preostalo = await UsageService.consume(user["user_id"], user.get("email", ""), "drafting")
        if req.predmet_id and isinstance(rezultat, dict) and rezultat.get("status") == "success" and rezultat.get("data"):
            asyncio.create_task(_stage_draft_for_review(user, req.predmet_id, req.vrsta, req.vrsta, rezultat["data"]))
        return _normalizuj_rezultat(rezultat, credits_remaining=max(preostalo, 0))
    except Exception as _exc:
        _sentry_capture(_exc)
        logger.exception("Greška u /api/nacrt")
        return _greska_odgovor(
            500,
            "Došlo je do greške na serveru. Pokušajte ponovo za nekoliko sekundi.",
        )


@router.post("/api/analiza")
@limiter.limit("10/minute")
async def analiza(req: AnalizaReq, request: Request, user: dict = Depends(PermissionService.require("drafting"))):
    """Analiza pravnog dokumenta."""
    qh = _q_hash(req.pitanje)
    logger.info("Analiza [uid=%.8s] [q=%s]", user["user_id"], qh)
    asyncio.create_task(_audit(user["user_id"], "analiza", qh))
    try:
        qh_analiza = _q_hash(_skini_pii(req.pitanje or req.tekst[:200]))
        t0 = _time.monotonic()
        rezultat = await _pokreni(ask_analiza, req.tekst, req.pitanje)
        latency_ms = int((_time.monotonic() - t0) * 1000)
        _al.log_response(
            endpoint="/api/analiza",
            query_hash=qh_analiza,
            response_text=rezultat.get("data", ""),
            latency_ms=latency_ms,
        )
        is_blocked = (rezultat.get("data") or "").startswith("[!] ANALIZA BLOKIRANA")
        if rezultat.get("status") == "success" and not is_blocked:
            preostalo = await UsageService.consume(user["user_id"], user.get("email", ""), "drafting")
        else:
            preostalo = await UsageService.balance(user["user_id"], user.get("email", ""))
        return _normalizuj_rezultat(rezultat, credits_remaining=max(preostalo, 0))
    except Exception as _exc:
        _sentry_capture(_exc)
        logger.exception("Neočekivana greška u /api/analiza")
        return _greska_odgovor(
            500,
            "Došlo je do greške na serveru. Pokušajte ponovo za nekoliko sekundi.",
        )


_FORMAT_INSTRUKCIJA = {
    "email": (
        "Format: EMAIL klijentu. Ton: profesionalan, formalan. "
        "Početak: 'Poštovani/a,' — zatim 4-5 rečenica — kraj: 'S poštovanjem, Vaš advokat'."
    ),
    "viber": (
        "Format: KRATKA VIBER PORUKA klijentu. Ton: neformalan, prijatan, bez 'Poštovani'. "
        "Maksimalno 3-4 kratke rečenice. Bez formalnog pozdrava na kraju."
    ),
    "pisano": (
        "Format: PISANO OBAVEŠTENJE klijentu. Ton: formalan, kao zvanično pismo. "
        "Koristiti 'Obaveštavamo Vas...' — 5-6 rečenica — zaključna napomena o konsultaciji."
    ),
}


@router.post("/api/sazmi")
@limiter.limit("10/minute")
async def sazmi(req: SazmiReq, request: Request, user: dict = Depends(PermissionService.require("drafting"))):
    """Generiše verziju odgovora na 'ljudskom' jeziku za klijenta (email/viber/pisano)."""
    from openai import OpenAI as _OAI
    try:
        fmt_instr = _FORMAT_INSTRUKCIJA.get(req.format, _FORMAT_INSTRUKCIJA["email"])
        klijent_prompt = (
            "Advokat ti šalje pravni odgovor koji treba da prepišeš za klijenta — laika koji ne zna pravo.\n"
            "PRAVILA TONA: Profesionalan, smiren, poverljiv. BEZ: latinštine, paragrafa, citata, 'čl.', 'Sl. glasnik', 'lex specialis'.\n\n"
            f"{fmt_instr}\n\n"
            "SADRŽAJ (obavezni elementi):\n"
            "  1. Šta znači situacija za klijenta — jasno i konkretno.\n"
            "  2. Šta je klijentov ključni dokaz ili sledeći korak — bez teorije.\n"
            "  3. Koji je rizik ako ne preduzme ništa (rok, zastarelost, gubitak prava).\n"
            "  4. Šta konkretno da uradi — imperativ.\n"
            "  5. Napomena: konsultacija sa advokatom pre svakog koraka.\n"
            "Počni direktno. Bez uvoda, bez 'Evo sažetka', bez zaglavlja."
        )
        client = _OAI(api_key=os.getenv("OPENAI_API_KEY"))
        # BUG FIX (2026-07-24): sinhroni SDK poziv unutar async def blokirao je
        # ceo event loop (i sve druge konkurentne zahteve na istom workeru) za
        # trajanje OpenAI poziva -- ostala 3 poziva u ovom fajlu (linije ~443,
        # 458, 507) su već ispravno await asyncio.to_thread(...)-ovana.
        resp = await asyncio.to_thread(
            _pozovi_drafting_api,
            client,
            model="gpt-4o-mini",
            temperature=0.3,
            max_tokens=600,
            messages=[
                {"role": "system", "content": klijent_prompt},
                {"role": "user",   "content": _skini_pii(req.odgovor[:4000])},
            ],
        )
        tekst = resp.choices[0].message.content.strip()
        await UsageService.consume(user["user_id"], user.get("email", ""), "drafting")
        return {"status": "ok", "sazetak": tekst, "format": req.format}
    except Exception as _exc:
        _sentry_capture(_exc)
        logger.exception("Greška u /api/sazmi")
        return _greska_odgovor(500, "Greška pri generisanju sažetka.")


@router.post("/api/feedback")
async def feedback(req: FeedbackReq, user: dict = Depends(get_current_user)):
    """
    Korisnik prijavljuje netačan ili nepotpun odgovor.
    NO-STORAGE POLICY (Basic API tier): čuvamo samo hash pitanja i tip — bez sadržaja.
    ZZPL čl. 5(1)(c) — minimizacija podataka.
    """
    try:
        qh = _q_hash(req.pitanje)
        await asyncio.to_thread(
            lambda: _get_supa().table("feedback").insert({
                "user_id": user["user_id"],
                "q_hash":  qh,
                "tip":     req.tip,
            }).execute()
        )
        logger.info("Feedback [uid=%.8s] tip=%s [q=%s]", user["user_id"], req.tip, qh)
        return {"status": "ok"}
    except Exception as _exc:
        _sentry_capture(_exc)
        logger.exception("Greška u /api/feedback")
        return {"status": "ok"}


@router.post("/api/podnesak")
@limiter.limit("5/minute")
async def podnesak(req: PodnesakReq, request: Request, user: dict = Depends(PermissionService.require("drafting"))):
    """
    Generiše nacrt sudskog podneska u dva koraka:
    1. Ekstrakcija entiteta iz slobodnog opisa (GPT-4o-mini, brzo)
    2. RAG + popunjavanje šablona (GPT-4o, precizno)
    """
    from openai import OpenAI as _OAI

    log_id = _q_hash(req.opis)
    logger.info("Podnesak [uid=%.8s] tip=%s sud=%s [q=%s]",
                user["user_id"], req.tip, req.sud_naziv or "-", log_id)

    oai = _OAI(api_key=os.getenv("OPENAI_API_KEY"))
    opis_api = _skini_pii(req.opis)

    # Pripremi kontekst sa sudom ako je korisnik izabrao
    sud_ctx = ""
    if req.sud_naziv:
        sud_ctx = (
            f"\n\nSUD (unapred određen — OBAVEZNO koristiti):\n"
            f"  Naziv: {req.sud_naziv}\n"
            f"  Adresa: {req.sud_adresa or 'N/A'}\n"
            f"Polja SUD_NAZIV i SUD_ADRESA popuniti ovim vrednostima."
        )

    ekstr_prompt = EKSTRAKCIONI_PROMPTOVI[req.tip]

    def _parse_json_safe(raw: str) -> dict:
        raw = raw.strip().removeprefix("```json").removeprefix("```").removesuffix("```").strip()
        try:
            return json.loads(raw)
        except json.JSONDecodeError:
            import re as _re
            m = _re.search(r'\{[\s\S]+\}', raw)
            if m:
                try:
                    return json.loads(m.group())
                except json.JSONDecodeError:
                    pass
        return {}

    try:
        ekstr_resp = await asyncio.to_thread(
            _pozovi_drafting_api,
            oai,
            model="gpt-4o-mini",
            temperature=0,
            max_tokens=900,
            messages=[
                {"role": "system", "content": ekstr_prompt},
                {"role": "user",   "content": f"OPIS SLUČAJA:\n{opis_api}{sud_ctx}"},
            ],
        )
        raw_json = (ekstr_resp.choices[0].message.content or "").strip()
        entiteti: dict = _parse_json_safe(raw_json)
        if not entiteti:
            logger.warning("Ekstrakcija vratila prazan JSON [q=%s] — retry sa gpt-4o", log_id)
            ekstr_resp2 = await asyncio.to_thread(
                _pozovi_drafting_api,
                oai,
                model="gpt-4o",
                temperature=0,
                max_tokens=900,
                messages=[
                    {"role": "system", "content": ekstr_prompt},
                    {"role": "user",   "content": f"OPIS SLUČAJA:\n{opis_api}\n\nVrati ISKLJUČIVO validan JSON objekat, bez ikakvog drugog teksta."},
                ],
            )
            entiteti = _parse_json_safe(ekstr_resp2.choices[0].message.content or "")
    except Exception as exc:
        _sentry_capture(exc)
        logger.warning("Ekstrakcija entiteta neuspešna [q=%s]: %s", log_id, exc)
        entiteti = {}

    rag_upit = f"{PODNESAK_TIPOVI[req.tip]}: {opis_api[:400]}"
    try:
        from app.services.retrieve import retrieve_documents
        # BUG FIX (2026-07-24): retrieve_documents vraća tuple[list[str], dict]
        # (docs, retrieval_meta), ne samu listu -- pre ove izmene je "docs"
        # bio ceo tuple, "\n\n".join(docs[:4]) je pucao sa TypeError na dict
        # elementu, spoljašnji except je to tiho gutao i kontekst je UVEK bio
        # prazan string. Svaki podnesak generisan pre ove izmene je pisan bez
        # ijednog stvarnog RAG zakonskog citata, uprkos promptu koji to
        # eksplicitno traži.
        docs, _retrieval_meta = await asyncio.to_thread(retrieve_documents, rag_upit, 5)
        kontekst = _izvori_kontekst(docs)
    except Exception as exc:
        _sentry_capture(exc)
        logger.warning("RAG neuspešan za podnesak [q=%s]: %s", log_id, exc)
        kontekst = ""

    vks_analiza = ""
    vks_kontekst_blok = ""
    if req.tip == "tuzba_naknada_stete":
        try:
            vks = vks_preporuci(entiteti)
            vks_kontekst_blok = f"\n\nVKS ORIJENTACIONI KRITERIJUMI:\n{vks['kontekst_tekst']}"
            vks_analiza = vks["analiza_tekst"]
        except Exception as exc:
            _sentry_capture(exc)
            logger.warning("VKS preporuka neuspešna [q=%s]: %s", log_id, exc)

    obog_prompt = OBOGACIVANJE_PROMPTOVI[req.tip]
    obog_user = (
        f"EKSTRAKTOVANI PODACI (JSON):\n{json.dumps(entiteti, ensure_ascii=False)}"
        f"{vks_kontekst_blok}\n\n"
        f"ZAKONSKI KONTEKST (RAG):\n{kontekst or 'Nije pronađen relevantan kontekst.'}"
    )
    try:
        obog_resp = await asyncio.to_thread(
            _pozovi_drafting_api,
            oai,
            model="gpt-4o",
            temperature=0,
            max_tokens=2500,
            messages=[
                {"role": "system", "content": obog_prompt},
                {"role": "user",   "content": obog_user},
            ],
        )
        raw_obog = (obog_resp.choices[0].message.content or "").strip()
        raw_obog = raw_obog.removeprefix("```json").removeprefix("```").removesuffix("```").strip()
        obogacivanje: dict = json.loads(raw_obog)
    except (json.JSONDecodeError, Exception) as exc:
        _sentry_capture(exc)
        logger.warning("Obogaćivanje neuspešno [q=%s]: %s", log_id, exc)
        obogacivanje = {}

    nacrt = popuni_sablon(req.tip, entiteti, obogacivanje, vks_analiza=vks_analiza)
    nacrt = await _critique_and_refine_draft(nacrt, kontekst, req.tip, log_id)

    await UsageService.consume(user["user_id"], user.get("email", ""), "drafting")

    if req.predmet_id and nacrt:
        asyncio.create_task(_stage_draft_for_review(user, req.predmet_id, req.tip, PODNESAK_TIPOVI[req.tip], nacrt))

    return {
        "status":  "success",
        "odgovor": nacrt,
        "tip":     req.tip,
        "naziv":   PODNESAK_TIPOVI[req.tip],
    }


@router.post("/api/nacrti/checklist")
@limiter.limit("20/minute")
async def nacrti_checklist(req: NacrtChecklistReq, request: Request, user: dict = Depends(PermissionService.require("drafting"))):
    """
    Faza 1 — Checklist analiza.

    Prima tip podneska i slobodan tekst činjenica.
    Vraća koji obavezni elementi nedostaju, sa objašnjenjem zašto su važni.
    blokira_nastavak=True ako nedostaje element kriticnosti "visoka".
    """
    from nacrti.checklist_engine import analiziraj_checklist

    log_id = _q_hash(req.cinjenice)
    logger.info("NacrtChecklist [uid=%.8s] tip=%s [q=%s]", user["user_id"], req.tip, log_id)

    try:
        rezultat = await asyncio.to_thread(analiziraj_checklist, req.tip, req.cinjenice)
    except KeyError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except RuntimeError as e:
        logger.error("NacrtChecklist GPT error [q=%s]: %s", log_id, e)
        raise HTTPException(status_code=502, detail="AI servis trenutno nedostupan. Pokušajte ponovo.")

    await UsageService.consume(user["user_id"], user.get("email", ""), "drafting")
    return rezultat


# ── DOCX Export ───────────────────────────────────────────────────────────────

class DocxExportReq(BaseModel):
    tekst:  str = Field(..., min_length=1, max_length=100_000)
    naslov: str = Field("Nacrt", max_length=200)
    tip:    str = Field("", max_length=100)


def _nacrt_to_docx(tekst: str, naslov: str, firma_info: dict | None = None) -> bytes:
    """Konvertuje nacrt tekst u .docx sa kancelarijskim headerom. Vraća bytes."""
    import io
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    if firma_info:
        naziv   = (firma_info.get("seller_naziv") or firma_info.get("naziv") or "").strip()
        adresa  = (firma_info.get("seller_adresa") or firma_info.get("adresa") or "").strip()
        mesto   = (firma_info.get("seller_mesto") or firma_info.get("mesto") or "").strip()
        pib     = (firma_info.get("seller_pib") or firma_info.get("pib") or "").strip()
        email   = (firma_info.get("email") or "").strip()

        if naziv:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(naziv)
            run.bold = True
            run.font.size = Pt(14)

        kontakt_delovi = [adresa, mesto, email]
        kontakt = " | ".join(d for d in kontakt_delovi if d)
        if kontakt:
            p = doc.add_paragraph(kontakt)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(10)
            p.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        if pib:
            p = doc.add_paragraph(f"PIB: {pib}")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = RGBColor(0x77, 0x77, 0x77)

        # Separator
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "999999")
        pBdr.append(bottom)
        pPr.append(pBdr)

    # Naslov
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(naslov.upper())
    run.bold = True
    run.font.size = Pt(13)
    doc.add_paragraph()

    # Sadržaj — markdown-aware parsing
    for linija in tekst.split("\n"):
        if linija.startswith("# "):
            doc.add_heading(linija[2:], level=1)
        elif linija.startswith("## "):
            doc.add_heading(linija[3:], level=2)
        elif linija.startswith("### "):
            doc.add_heading(linija[4:], level=3)
        elif linija.strip() == "":
            doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            parts = linija.split("**")
            for idx, part in enumerate(parts):
                if part:
                    run = p.add_run(part)
                    run.bold = (idx % 2 == 1)
                    run.font.size = Pt(11)

    # Footer
    from datetime import date as _date
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = fp.add_run(f"Generisano: {_date.today().strftime('%d.%m.%Y.')} | Vindex AI")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


@router.post("/api/nacrti/export/docx")
@limiter.limit("30/minute")
async def export_nacrt_docx(
    req: DocxExportReq,
    request: Request,
    user: dict = Depends(get_current_user),
):
    """Preuzmi nacrt kao .docx fajl sa kancelarijskim headerom."""
    uid = user["user_id"]
    supa = _get_supa()

    # Dohvati firma/SEF info za header
    firma_info = None
    try:
        r = await asyncio.to_thread(
            lambda: supa.table("sef_podesavanja")
                .select("seller_pib,seller_naziv,seller_adresa,seller_mesto")
                .eq("user_id", uid)
                .maybe_single()
                .execute()
        )
        if r and r.data:
            firma_info = r.data
    except Exception as _exc:
        _sentry_capture(_exc)
        pass

    naslov = req.naslov or req.tip or "Nacrt"

    try:
        docx_bytes = await asyncio.to_thread(_nacrt_to_docx, req.tekst, naslov, firma_info)
    except Exception as e:
        logger.error("DOCX export greška: %s", e)
        raise HTTPException(status_code=500, detail="Greška pri generisanju DOCX fajla.")

    safe_name = "".join(c if c.isalnum() or c in " _-" else "_" for c in naslov)[:40]
    filename = f"{safe_name}.docx"
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ─── Institutional Memory V2 (2026-07-26) — Staging Memory review endpoints ──

@router.get("/api/staging/predmet/{predmet_id}")
@limiter.limit("30/minute")
async def staging_list_za_predmet(predmet_id: str, request: Request, user: dict = Depends(get_current_user)):
    """Lista nacrta koji čekaju (ili su prošli) advokatsku overu za dati predmet."""
    supa = _get_supa()
    r = await asyncio.to_thread(
        lambda: supa.table("staging_memory")
            .select("id,tip,naziv,confidence_score,is_lawyer_approved,status,pinecone_indexed,created_at")
            .eq("predmet_id", predmet_id)
            .eq("user_id", user["user_id"])
            .order("created_at", desc=True)
            .execute()
    )
    return {"stavke": r.data or []}


@router.post("/api/staging/{staging_id}/approve")
@limiter.limit("20/minute")
async def staging_approve(staging_id: str, request: Request, user: dict = Depends(get_current_user)):
    """Eksplicitna advokatska potvrda (STUB 1.3). NUŽAN ali NE DOVOLJAN uslov
    za ulazak u Pinecone -- promocija se dešava SAMO ako je i
    confidence_score >= 0.85 (v. _promote_staged_draft_to_pinecone)."""
    supa = _get_supa()
    row_res = await asyncio.to_thread(
        lambda: supa.table("staging_memory").select("*").eq("id", staging_id).eq("user_id", user["user_id"]).maybe_single().execute()
    )
    if not row_res.data:
        raise HTTPException(status_code=404, detail="Nacrt na čekanju nije pronađen.")
    staging_row = row_res.data

    from shared.vector_origin import now_iso
    update_fields = {
        "is_lawyer_approved": True,
        "approved_by": user["user_id"],
        "approved_at": now_iso(),
        "status": "approved",
    }

    promoted = False
    if float(staging_row.get("confidence_score") or 0) >= _APPROVAL_CONFIDENCE_THRESHOLD:
        try:
            promoted = await _promote_staged_draft_to_pinecone(supa, {**staging_row, **update_fields})
        except Exception as exc:
            _sentry_capture(exc)
            logger.warning("[STAGING_APPROVE] promocija neuspešna staging_id=%s: %s", staging_id, exc)
            promoted = False

    update_fields["pinecone_indexed"] = promoted
    await asyncio.to_thread(
        lambda: supa.table("staging_memory").update(update_fields).eq("id", staging_id).execute()
    )

    if promoted:
        return {"status": "approved", "indexed": True, "poruka": "Nacrt odobren i dodat u bazu znanja kancelarije."}
    return {
        "status": "approved", "indexed": False,
        "poruka": (
            "Nacrt je označen kao odobren, ali confidence_score "
            f"({staging_row.get('confidence_score')}) je ispod praga {_APPROVAL_CONFIDENCE_THRESHOLD} "
            "za automatski ulazak u bazu znanja kancelarije."
        ),
    }


@router.post("/api/staging/{staging_id}/reject")
@limiter.limit("20/minute")
async def staging_reject(staging_id: str, request: Request, user: dict = Depends(get_current_user)):
    """Advokat odbija nacrt -- nikad ne ulazi u Pinecone, ostaje samo istorijski
    zapis u staging_memory (audit trail — zašto je nešto odbijeno)."""
    supa = _get_supa()
    row_res = await asyncio.to_thread(
        lambda: supa.table("staging_memory").select("id").eq("id", staging_id).eq("user_id", user["user_id"]).maybe_single().execute()
    )
    if not row_res.data:
        raise HTTPException(status_code=404, detail="Nacrt na čekanju nije pronađen.")

    await asyncio.to_thread(
        lambda: supa.table("staging_memory").update({"status": "rejected", "is_lawyer_approved": False}).eq("id", staging_id).execute()
    )
    return {"status": "rejected", "indexed": False}
