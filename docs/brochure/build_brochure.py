# -*- coding: utf-8 -*-
"""
Generator brošure Vindex AI — PDF + DOCX + MD iz jednog izvora sadržaja.

SVAKA TVRDNJA U OVOM FAJLU JE PROVERENA IZ REPOZITORIJUMA.
Verifikaciona tabela je u docs/brochure/VERIFICATION.md — svaka stavka ima
izvor (fajl/mehanizam) i status. Ono što nije moglo da se dokaže NIJE u brošuri.

Vizuelni jezik preuzet sa postojeće landing strane (static/): tamna podloga
#010308, akcenat #00d4ff, tekst #e6edf3. Oštri uglovi, bez gradijenata i bez
generičke AI ikonografije.

Pokretanje:  python docs/brochure/build_brochure.py
Zavisnosti:  reportlab, python-docx  (obe već prisutne u okruženju)
"""
from __future__ import annotations

import os
from pathlib import Path

OUT = Path(__file__).parent
BG, ACCENT, TEXT, MUTED, RULE = "#010308", "#00d4ff", "#e6edf3", "#8b98a5", "#1c2530"

# ─────────────────────────────────────────────────────────────────────────────
# SADRŽAJ — jedan izvor istine za sva tri formata
# ─────────────────────────────────────────────────────────────────────────────

PAGES: list[dict] = [
    {
        "kind": "cover",
        "title": "VINDEX AI",
        "subtitle": "AI infrastruktura za složen profesionalni rad",
        "body": [
            "Vindex AI je platforma koja povezuje dokumente, kontekst predmeta, "
            "znanje i AI modele u jedan radni sloj — sa proverljivim tragom o "
            "tome odakle svaki podatak dolazi.",
            "Prvo tržište: advokatske kancelarije.",
        ],
        "foot": "Poverljivo — pripremljeno za poslovne sagovornike",
    },
    {
        "kind": "page",
        "title": "PROBLEM",
        "lead": "Vreme profesionalca troši se na obradu informacija, ne na stručnu procenu.",
        "body": [
            "Rad na složenom predmetu znači rad sa velikim brojem dokumenata koji "
            "nisu međusobno povezani. Činjenice su razbacane po podnescima, "
            "prilozima i prepisci. Rokovi zavise od datuma koji se nalaze u tekstu, "
            "a ne u kalendaru.",
            "Kada predmet traje mesecima, kontekst se gubi. Onaj ko se vrati na "
            "predmet posle nekoliko nedelja mora ponovo da rekonstruiše šta je "
            "utvrđeno, šta je sporno i šta sledi.",
            "Cena propuštene informacije je nesrazmerno visoka, a provera je skupa "
            "jer zahteva ponovno čitanje.",
        ],
        "note": "Ovaj opis proizlazi iz problema koji platforma rešava. Ne sadrži "
                "statistiku koju nismo sami izmerili.",
    },
    {
        "kind": "page",
        "title": "PRISTUP",
        "lead": "Vindex nije ćaskanje sa modelom. Vindex je sloj između predmeta i modela.",
        "body": [
            "Razlika je u tome šta model dobija. Kod običnog AI alata korisnik "
            "sam sastavlja pitanje i sam bira šta će zalepiti u njega. Kod Vindexa "
            "sistem održava strukturisan prikaz predmeta i taj prikaz prosleđuje "
            "modelu.",
            "Svako polje tog prikaza nosi oznaku porekla — iz kog dokumenta potiče, "
            "ko ga je uneo i kada se osvežava. Zahvaljujući tome moguće je "
            "razlikovati utvrđenu činjenicu od zaključka koji je neko izveo.",
        ],
        "diagram": [
            ("DOKUMENTI", "unos, prepoznavanje teksta, klasifikacija"),
            ("KONTEKST PREDMETA", "činjenice, rokovi, dokazi, poreklo podatka"),
            ("AI SLOJ", "model dobija strukturisan prikaz, ne sirov tekst"),
            ("REZULTAT", "analiza, nacrt, upozorenje — uz trag do izvora"),
            ("TRAG", "nepromenljiv zapis ko je šta pokrenuo i kada"),
        ],
    },
    {
        "kind": "page",
        "title": "ŠTA VINDEX DANAS RADI",
        "lead": "Ispod je samo ono što je implementirano i proverljivo u sistemu.",
        "groups": [
            ("OBRADA DOKUMENATA",
             "Unos većeg broja dokumenata odjednom, prepoznavanje teksta iz "
             "skeniranih fajlova, automatska klasifikacija i povezivanje sa "
             "predmetom. Original se čuva."),
            ("RAZUMEVANJE PREDMETA",
             "Strukturisan prikaz predmeta sa oznakom porekla za svako polje. "
             "Hronologija, stranke, dokazi i rokovi vode se kao povezani podaci, "
             "ne kao slobodan tekst."),
            ("PRETRAGA I ZNANJE",
             "Semantička pretraga po sadržaju dokumenata i po bazi propisa i "
             "prakse, umesto pretrage po ključnoj reči."),
            ("ANALIZA",
             "Prepoznavanje rokova i obaveza, uočavanje protivrečnosti između "
             "dokumenata, procena rizika predmeta."),
            ("IZRADA PODNESAKA",
             "Generisanje nacrta vezanog za konkretan predmet i njegove dokumente."),
            ("EVIDENCIJA I BEZBEDNOST",
             "Nepromenljiv zapis radnji, poreklo svakog AI odgovora, razdvajanje "
             "podataka između korisnika i kancelarija."),
        ],
    },
    {
        "kind": "page",
        "title": "ZAŠTO PRVO PRAVO",
        "lead": "Advokatura je prvo tržište jer je najzahtevnija sredina za proveru, "
                "a ne zato što je jedina.",
        "body": [
            "Pravni rad ima svojstva koja se retko sreću zajedno: veliki obim "
            "dokumenata, složene veze između njih, potrebu da se iz teksta izvuku "
            "činjenice i datumi, obavezujuće rokove, visoku osetljivost podataka i "
            "zahtev da se svaka tvrdnja može vratiti na izvor.",
            "Sistem koji izdrži takav rad primenljiv je i tamo gde su zahtevi blaži. "
            "Zato je advokatura izabrana kao okruženje za proveru — ne kao granica "
            "poslovnog dometa.",
        ],
    },
    {
        "kind": "page",
        "title": "OD PRAVA KA DRUGIM DELATNOSTIMA",
        "lead": "Sposobnosti platforme nisu vezane za pravnu materiju.",
        "expansion": [
            "ADVOKATURA — prvo tržište, sredina za proveru",
            "NOTARIJAT — velike količine isprava, strogi formalni zahtevi",
            "BANKARSTVO — dokumentacija, provera podataka, procena rizika",
            "OSIGURANJE — obrada odštetnih zahteva i prateće dokumentacije",
            "KORPORATIVNI PRAVNI POSLOVI — ugovori, obaveze, rokovi",
            "KONSALTING I REVIZIJA — rad zasnovan na dokumentima i analizi",
        ],
        "note": "Navedene delatnosti su mogućnosti proširenja. Vindex u njima "
                "trenutno nema korisnike niti sprovedenu proveru tržišta.",
    },
    {
        "kind": "page",
        "title": "BEZBEDNOST I POVERENJE",
        "lead": "Bezbednost je deo proizvoda, ne dodatak.",
        "groups": [
            ("RAZDVAJANJE PODATAKA",
             "Pripadnost svakog zapisa proverava se u samoj operaciji nad bazom, "
             "a ne samo pre nje. Ovaj mehanizam je bio predmet više nezavisnih "
             "unutrašnjih provera."),
            ("NEPROMENLJIVA EVIDENCIJA",
             "Zapisi o radnjama vezani su u lanac zaštićen kriptografskim otiskom "
             "i ograničenjima na nivou baze koja sprečavaju naknadnu izmenu ili "
             "brisanje."),
            ("POREKLO AI ODGOVORA",
             "Za svaki AI poziv beleži se koji je model korišćen, u okviru kog "
             "predmeta i sa kojim identifikatorom zahteva. Sadržaj upita i "
             "odgovora se u tu evidenciju ne upisuje."),
            ("KONTROLA PRISTUPA",
             "Prava unutar kancelarije razdvojena su po ulogama; administrativne "
             "radnje odbijaju se korisniku bez odgovarajućeg ovlašćenja."),
            ("ZAŠTITA UPITA",
             "Ulazni sadržaj prolazi kroz proveru pre nego što se prosledi modelu."),
        ],
        "note": "Ne tvrdimo potpunu bezbednost niti posedovanje sertifikata. "
                "Navedeni su mehanizmi koji postoje u sistemu.",
    },
    {
        "kind": "page",
        "title": "ARHITEKTURA AI SLOJA",
        "lead": "Model je komponenta koju platforma koristi, a ne sam proizvod.",
        "body": [
            "Vindex je projektovan tako da AI model bude zamenljiva komponenta. "
            "Platforma održava kontekst predmeta, pravila i evidenciju; model "
            "obavlja pojedinačan zadatak nad onim što mu platforma prosledi.",
            "U sistemu postoji sloj koji odvaja Vindex od pojedinačnog dobavljača "
            "AI modela: jedinstven oblik zahteva i odgovora, kontrola pre poziva, "
            "normalizovane greške i evidencija posle poziva. Time se izbegava "
            "vezanost proizvoda za jednog dobavljača.",
        ],
        "status": [
            ("U PRODUKCIJI", "Jedan dobavljač AI modela opslužuje sve postojeće "
                             "funkcije."),
            ("IMPLEMENTIRANO, NIJE U PRODUKCIONOM TOKU",
             "Sloj za rad sa više dobavljača, sa pripremljenim priključcima za "
             "tri različita dobavljača. Nijedna postojeća funkcija još ne ide "
             "kroz taj sloj."),
            ("PLANIRANO", "Unakrsna provera odgovora između dva modela."),
        ],
    },
    {
        "kind": "page",
        "title": "TRENUTNO STANJE",
        "lead": "Proizvod je u završnoj pripremi pred testiranje sa prvim korisnicima.",
        "status": [
            ("RAZVOJ", "Funkcionalna celina je implementirana i pokrivena "
                       "automatizovanim testovima koji se izvršavaju nad svakom "
                       "izmenom."),
            ("PROVERA TRŽIŠTA", "U toku su razgovori sa više advokata. Ostvaren je "
                                "i kontakt sa Stefanom Gojkovićem, sudijskim "
                                "pomoćnikom."),
            ("TESTIRANJE", "Priprema zatvorenog testiranja sa ograničenim brojem "
                           "korisnika."),
            ("PRISUTNOST NA MREŽI", "Postoji uvodna stranica. Potpun sajt je u planu."),
        ],
        "note": "Sagovornici navedeni iznad nisu korisnici, klijenti ni poslovni "
                "partneri. Reč je o ranim razgovorima radi provere pretpostavki.",
    },
    {
        "kind": "page",
        "title": "VIZIJA",
        "lead": "Od pravnog proizvoda ka infrastrukturi za profesionalni rad.",
        "body": [
            "Dugoročni cilj nije da Vindex bude još jedan alat koji odgovara na "
            "pitanja. Cilj je da bude sloj koji drži uređen prikaz stvarnosti "
            "jedne organizacije — dokumenata, obaveza, odluka i njihovog porekla — "
            "i da taj prikaz stavlja na raspolaganje AI modelima, ma koji model "
            "sutra bio najbolji.",
            "Vrednost u tom slučaju ne leži u modelu, koji se menja, nego u "
            "uređenom kontekstu i proverljivom tragu, koji ostaju.",
        ],
    },
    {
        "kind": "closing",
        "title": "SLEDEĆI KORAK",
        "body": [
            "Vindex AI trenutno traži ograničen broj advokatskih kancelarija "
            "spremnih da učestvuju u zatvorenom testiranju.",
            "Za razgovor o učešću, partnerstvu ili proveri primenljivosti u drugoj "
            "delatnosti, obratite se osnivaču putem kanala kojim ste primili ovaj "
            "dokument.",
        ],
        "foot": "Ovaj dokument opisuje stanje proizvoda u trenutku izrade. "
                "Ne sadrži obećanja o učinku niti garancije rezultata.",
    },
]


# ─────────────────────────────────────────────────────────────────────────────
# MARKDOWN
# ─────────────────────────────────────────────────────────────────────────────

def build_md() -> str:
    L: list[str] = ["# VINDEX AI", "", "**AI infrastruktura za složen profesionalni rad**", ""]
    for p in PAGES:
        if p["kind"] == "cover":
            L += [f"> {p['subtitle']}", ""]
            L += [b for b in p["body"]] + [""]
            continue
        L += [f"## {p['title']}", ""]
        if p.get("lead"):
            L += [f"*{p['lead']}*", ""]
        for b in p.get("body", []):
            L += [b, ""]
        for name, desc in p.get("groups", []):
            L += [f"**{name}** — {desc}", ""]
        for name, desc in p.get("diagram", []):
            L += [f"- **{name}** — {desc}"]
        if p.get("diagram"):
            L += [""]
        for item in p.get("expansion", []):
            L += [f"- {item}"]
        if p.get("expansion"):
            L += [""]
        for name, desc in p.get("status", []):
            L += [f"**{name}** — {desc}", ""]
        if p.get("note"):
            L += [f"> {p['note']}", ""]
        if p.get("foot"):
            L += [f"_{p['foot']}_", ""]
    return "\n".join(L)


# ─────────────────────────────────────────────────────────────────────────────
# PDF (reportlab)
# ─────────────────────────────────────────────────────────────────────────────

def _register_fonts() -> tuple[str, str, str]:
    """Ugrađuje TrueType font sa punom srpskom latinicom (č ć ž š đ Č Ć Ž Š Đ).

    Ugrađeni PDF font Helvetica koristi WinAnsi kodiranje koje ta slova NEMA --
    otuda kockice u ranijoj verziji. Ovde se registruje pravi TTF; ako nijedan
    kandidat nije dostupan, funkcija diže grešku umesto da tiho vrati Helveticu
    i ponovo proizvede neispravan dokument.
    """
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    kandidati = [
        # (porodica, regular, bold, italic) -- redosled je i redosled prioriteta
        ("Calibri", "calibri.ttf", "calibrib.ttf", "calibrii.ttf"),
        ("Arial",   "arial.ttf",   "arialbd.ttf",  "ariali.ttf"),
        ("DejaVuSans", "DejaVuSans.ttf", "DejaVuSans-Bold.ttf", "DejaVuSans-Oblique.ttf"),
    ]
    dirs = [Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts",
            Path("/usr/share/fonts/truetype/dejavu"),
            Path("/Library/Fonts")]

    for fam, reg, bold, ital in kandidati:
        for d in dirs:
            p_reg = d / reg
            if not p_reg.exists():
                continue
            p_bold = d / bold if (d / bold).exists() else p_reg
            p_ital = d / ital if (d / ital).exists() else p_reg
            pdfmetrics.registerFont(TTFont(fam, str(p_reg)))
            pdfmetrics.registerFont(TTFont(fam + "-B", str(p_bold)))
            pdfmetrics.registerFont(TTFont(fam + "-I", str(p_ital)))
            return fam, fam + "-B", fam + "-I"

    raise RuntimeError(
        "Nije pronađen nijedan TrueType font sa srpskom latinicom. "
        "PDF se NE generiše sa Helveticom jer bi č/ć/ž/š/đ ostali neispravni."
    )


def build_pdf(path: Path) -> int:
    from reportlab.lib.colors import HexColor
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.pdfgen import canvas as rl_canvas

    F, F_B, F_I = _register_fonts()
    W, H = A4
    M = 22 * mm
    c = rl_canvas.Canvas(str(path), pagesize=A4)
    pages = 0

    def bg():
        c.setFillColor(HexColor(BG)); c.rect(0, 0, W, H, stroke=0, fill=1)

    def wrap(text: str, font: str, size: float, maxw: float) -> list[str]:
        c.setFont(font, size)
        out, line = [], ""
        for word in text.split():
            t = (line + " " + word).strip()
            if c.stringWidth(t, font, size) <= maxw:
                line = t
            else:
                out.append(line); line = word
        if line:
            out.append(line)
        return out

    def para(text, y, font=None, size=14, lead=21, color=TEXT, maxw=None):
        font = font or F
        maxw = maxw or (W - 2 * M)
        c.setFillColor(HexColor(color))
        for ln in wrap(text, font, size, maxw):
            c.setFont(font, size); c.drawString(M, y, ln); y -= lead
        return y

    # ── naslovna
    p = PAGES[0]; bg(); pages += 1
    c.setFillColor(HexColor(ACCENT)); c.rect(M, H - 62 * mm, 34 * mm, 1.4, stroke=0, fill=1)
    c.setFillColor(HexColor(TEXT)); c.setFont(F_B, 44)
    c.drawString(M, H - 82 * mm, p["title"])
    c.setFillColor(HexColor(ACCENT)); c.setFont(F, 15.5)
    c.drawString(M, H - 94 * mm, p["subtitle"])
    y = H - 118 * mm
    for b in p["body"]:
        y = para(b, y, size=14, lead=21.5, color=MUTED, maxw=W - 2 * M - 30 * mm)
        y -= 5
    c.setFillColor(HexColor(RULE)); c.rect(M, 34 * mm, W - 2 * M, 0.7, stroke=0, fill=1)
    c.setFillColor(HexColor(MUTED)); c.setFont(F, 9.5)
    c.drawString(M, 27 * mm, p["foot"])
    c.showPage()

    # ── unutrašnje strane
    for idx, p in enumerate(PAGES[1:], start=2):
        bg(); pages += 1
        c.setFillColor(HexColor(ACCENT)); c.rect(M, H - 30 * mm, 14 * mm, 1.2, stroke=0, fill=1)
        c.setFillColor(HexColor(TEXT)); c.setFont(F_B, 23)
        c.drawString(M, H - 41 * mm, p["title"])
        y = H - 54 * mm
        if p.get("lead"):
            y = para(p["lead"], y, F_I, 15, 22, ACCENT)
            y -= 9
        for b in p.get("body", []):
            y = para(b, y); y -= 8
        for name, desc in p.get("groups", []):
            c.setFillColor(HexColor(ACCENT)); c.setFont(F_B, 12.6)
            c.drawString(M, y, name); y -= 14
            y = para(desc, y, size=13.4, lead=19.5, color=MUTED); y -= 11
        for name, desc in p.get("diagram", []):
            c.setFillColor(HexColor(RULE)); c.rect(M, y - 5, W - 2 * M, 0.6, stroke=0, fill=1)
            c.setFillColor(HexColor(ACCENT)); c.setFont(F_B, 12.4)
            c.drawString(M, y + 5, name)
            c.setFillColor(HexColor(MUTED)); c.setFont(F, 12.4)
            c.drawRightString(W - M, y + 5, desc); y -= 29
        for item in p.get("expansion", []):
            c.setFillColor(HexColor(ACCENT)); c.circle(M + 2, y + 4, 1.8, stroke=0, fill=1)
            c.setFillColor(HexColor(TEXT)); c.setFont(F, 13.6)
            c.drawString(M + 10, y, item); y -= 24
        for name, desc in p.get("status", []):
            c.setFillColor(HexColor(ACCENT)); c.setFont(F_B, 12)
            c.drawString(M, y, name); y -= 13
            y = para(desc, y, size=13.4, lead=19.5, color=MUTED); y -= 11
        if p.get("note"):
            y -= 6
            c.setFillColor(HexColor(RULE)); c.rect(M, y - 2, 2, -42, stroke=0, fill=1)
            para(p["note"], y - 12, F_I, 11.6, 16.5, MUTED, W - 2 * M - 14)
        if p.get("foot"):
            c.setFillColor(HexColor(MUTED)); c.setFont(F_I, 9.6)
            c.drawString(M, 27 * mm, p["foot"])
        c.setFillColor(HexColor(MUTED)); c.setFont(F, 9.2)
        c.drawRightString(W - M, 18 * mm, f"VINDEX AI  ·  {idx}")
        if y < 34 * mm:
            print(f"  UPOZORENJE: strana {idx} ({p['title']}) prelazi dno (y={y/mm:.0f}mm)")
        c.showPage()

    c.save()
    return pages


# ─────────────────────────────────────────────────────────────────────────────
# DOCX
# ─────────────────────────────────────────────────────────────────────────────

def build_docx(path: Path) -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor

    d = Document()
    for p in PAGES:
        if p["kind"] == "cover":
            h = d.add_heading(p["title"], level=0)
            d.add_paragraph(p["subtitle"])
            for b in p["body"]:
                d.add_paragraph(b)
            d.add_page_break()
            continue
        d.add_heading(p["title"], level=1)
        if p.get("lead"):
            it = d.add_paragraph().add_run(p["lead"]); it.italic = True
        for b in p.get("body", []):
            d.add_paragraph(b)
        for name, desc in p.get("groups", []) + p.get("status", []):
            par = d.add_paragraph(); par.add_run(name + " — ").bold = True
            par.add_run(desc)
        for name, desc in p.get("diagram", []):
            par = d.add_paragraph(style="List Bullet")
            par.add_run(name + " — ").bold = True
            par.add_run(desc)
        for item in p.get("expansion", []):
            d.add_paragraph(item, style="List Bullet")
        if p.get("note"):
            n = d.add_paragraph().add_run("Napomena: " + p["note"]); n.italic = True
        if p.get("foot"):
            f = d.add_paragraph().add_run(p["foot"]); f.italic = True
        d.add_page_break()
    d.save(str(path))


if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    (OUT / "Vindex_AI_Brochure.md").write_text(build_md(), encoding="utf-8")
    n = build_pdf(OUT / "Vindex_AI_Brochure.pdf")
    build_docx(OUT / "Vindex_AI_Brochure_Source.docx")
    print(f"PDF strana: {n}")
    for f in sorted(OUT.glob("Vindex_AI_Brochure*")):
        print(f"  {f.name}  {f.stat().st_size:,} B")
